#!/usr/bin/env python3
"""
简单的Word文档创建脚本
使用python-docx库（如果可用）
"""

import os
import sys
import subprocess

def check_docx_module():
    """检查python-docx模块是否可用"""
    try:
        import docx
        return True
    except ImportError:
        return False

def install_docx_module():
    """尝试安装python-docx模块"""
    print("正在安装python-docx模块...")
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        return True
    except Exception as e:
        print(f"安装失败: {e}")
        return False

def create_docx():
    """创建Word文档"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 读取文本内容
        with open('交易咨询专员设置方案.txt', 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 创建文档
        doc = Document()
        
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)
        
        # 添加标题
        title = doc.add_paragraph()
        title_run = title.add_run('关于在业务单位设置交易咨询专员协助研究所开展交易咨询项目的请示')
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加副标题
        doc.add_paragraph()
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run('申万期货研究所')
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.bold = True
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_para = doc.add_paragraph()
        date_run = date_para.add_run('2026年3月')
        date_run.font.size = Pt(12)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # 添加正文内容
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
                
            # 简单处理标题
            if line.startswith('一、') or line.startswith('二、') or line.startswith('三、'):
                para = doc.add_paragraph()
                run = para.add_run(line)
                run.font.size = Pt(14)
                run.font.bold = True
                
            elif line.startswith('1.') or line.startswith('2.') or line.startswith('3.'):
                para = doc.add_paragraph(style='List Number')
                run = para.add_run(line[2:].strip())
                
            elif line.startswith('- ') or line.startswith('* '):
                para = doc.add_paragraph(style='List Bullet')
                run = para.add_run(line[2:].strip())
                
            else:
                para = doc.add_paragraph(line)
        
        # 保存文档
        doc.save('交易咨询专员设置方案.docx')
        print("Word文档创建成功!")
        return True
        
    except Exception as e:
        print(f"创建文档时出错: {e}")
        return False

def main():
    print("开始创建交易咨询方案Word文档...")
    
    # 检查模块
    if not check_docx_module():
        print("python-docx模块未安装")
        response = input("是否尝试安装? (y/n): ")
        if response.lower() == 'y':
            if install_docx_module():
                print("模块安装成功!")
            else:
                print("模块安装失败，无法创建Word文档")
                return False
        else:
            print("未安装模块，无法创建Word文档")
            return False
    
    # 创建文档
    return create_docx()

if __name__ == "__main__":
    success = main()
    if success:
        print("文档已保存为: 交易咨询专员设置方案.docx")
    else:
        print("文档创建失败")