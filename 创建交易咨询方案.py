#!/usr/bin/env python3
"""
创建交易咨询专员设置方案Word文档
"""

import os
import sys
import json
from pathlib import Path

# 添加docx技能路径
docx_skill_path = Path(__file__).parent / "active_skills" / "docx"
scripts_path = docx_skill_path / "scripts"
if scripts_path.exists():
    sys.path.insert(0, str(docx_skill_path))
    sys.path.insert(0, str(scripts_path))

# 尝试导入docx-js相关模块
try:
    from docx import Document, Packer, Paragraph, TextRun, Table, TableRow, TableCell
    from docx import Header, Footer, PageNumber, PageBreak, ExternalHyperlink
    from docx import TableOfContents, HeadingLevel
    from docx.enum.text import AlignmentType
    from docx.enum.table import WidthType
    from docx.enum.section import PageOrientation
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("警告: docx模块未安装，尝试使用替代方法")

def create_docx_with_docxjs(content_md, output_path):
    """使用docx-js库创建Word文档"""
    try:
        # 读取markdown内容
        with open(content_md, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # 解析内容
        lines = md_content.split('\n')
        
        # 创建文档
        doc = Document()
        
        # 设置页面大小和边距
        section = doc.sections[0]
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        
        # 添加标题
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run('关于在业务单位设置交易咨询专员协助研究所开展交易咨询项目的请示')
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_paragraph.alignment = AlignmentType.CENTER
        
        doc.add_paragraph()  # 空行
        
        # 添加副标题
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run('申万期货研究所')
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.bold = True
        subtitle.alignment = AlignmentType.CENTER
        
        date_para = doc.add_paragraph()
        date_run = date_para.add_run('2026年3月')
        date_run.font.size = Pt(12)
        date_para.alignment = AlignmentType.CENTER
        
        doc.add_page_break()
        
        # 解析markdown并添加内容
        in_table = False
        current_table = None
        current_row = None
        
        for line in lines:
            line = line.strip()
            
            if not line:
                doc.add_paragraph()
                continue
                
            # 处理标题
            if line.startswith('# '):
                title = line[2:].strip()
                para = doc.add_paragraph()
                run = para.add_run(title)
                run.font.size = Pt(16)
                run.font.bold = True
                para.alignment = AlignmentType.CENTER
                
            elif line.startswith('## '):
                title = line[3:].strip()
                para = doc.add_paragraph()
                run = para.add_run(title)
                run.font.size = Pt(14)
                run.font.bold = True
                
            elif line.startswith('### '):
                title = line[4:].strip()
                para = doc.add_paragraph()
                run = para.add_run(title)
                run.font.size = Pt(12)
                run.font.bold = True
                
            # 处理表格
            elif '|' in line and ('---' in line or '----' in line):
                # 表格标题行
                in_table = True
                headers = [h.strip() for h in line.split('|')[1:-1]]
                
            elif '|' in line and in_table:
                # 表格数据行
                cells = [c.strip() for c in line.split('|')[1:-1]]
                
                if not current_table:
                    current_table = doc.add_table(rows=1, cols=len(headers))
                    current_table.style = 'Light Grid'
                    
                    # 添加表头
                    header_cells = current_table.rows[0].cells
                    for i, header in enumerate(headers):
                        header_cells[i].text = header
                        header_cells[i].paragraphs[0].runs[0].font.bold = True
                        
                    # 添加数据行
                    row = current_table.add_row()
                    for i, cell in enumerate(cells):
                        row.cells[i].text = cell
                        
                else:
                    row = current_table.add_row()
                    for i, cell in enumerate(cells):
                        row.cells[i].text = cell
                        
            elif in_table and not '|' in line:
                # 表格结束
                in_table = False
                current_table = None
                doc.add_paragraph()
                
            # 处理列表项
            elif line.startswith('- ') or line.startswith('* '):
                item = line[2:].strip()
                para = doc.add_paragraph(style='List Bullet')
                para.add_run(item)
                
            elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
                item = line[3:].strip()
                para = doc.add_paragraph(style='List Number')
                para.add_run(item)
                
            # 普通段落
            else:
                para = doc.add_paragraph(line)
                
        # 保存文档
        doc.save(output_path)
        print(f"文档已保存到: {output_path}")
        return True
        
    except Exception as e:
        print(f"创建Word文档时出错: {e}")
        return False

def create_docx_alternative(content_md, output_path):
    """使用替代方法创建文档"""
    try:
        # 使用pandoc转换markdown到docx
        import subprocess
        
        # 检查pandoc是否可用
        try:
            subprocess.run(['pandoc', '--version'], capture_output=True, check=True)
            has_pandoc = True
        except (subprocess.CalledProcessError, FileNotFoundError):
            has_pandoc = False
            
        if has_pandoc:
            cmd = ['pandoc', content_md, '-o', output_path, '--reference-doc', 'template.docx']
            result = subprocess.run(cmd, capture_output=True, text=True)
            if result.returncode == 0:
                print(f"使用pandoc成功创建文档: {output_path}")
                return True
            else:
                print(f"pandoc转换失败: {result.stderr}")
        else:
            print("pandoc未安装，无法创建Word文档")
            
        # 如果pandoc不可用，创建纯文本文件
        print("创建纯文本版本...")
        with open(content_md, 'r', encoding='utf-8') as f:
            content = f.read()
            
        txt_output = output_path.replace('.docx', '.txt')
        with open(txt_output, 'w', encoding='utf-8') as f:
            f.write(content)
            
        print(f"纯文本版本已保存到: {txt_output}")
        return False
        
    except Exception as e:
        print(f"替代方法创建文档失败: {e}")
        return False

def main():
    # 输入文件和输出文件
    content_md = "交易咨询专员设置方案初稿.md"
    output_docx = "交易咨询专员设置方案.docx"
    
    if not os.path.exists(content_md):
        print(f"错误: 找不到输入文件 {content_md}")
        return False
    
    print(f"从 {content_md} 创建Word文档...")
    
    # 尝试使用docx-js
    if HAS_DOCX:
        print("使用docx-js库创建文档...")
        if create_docx_with_docxjs(content_md, output_docx):
            return True
        else:
            print("docx-js创建失败，尝试替代方法...")
    
    # 使用替代方法
    return create_docx_alternative(content_md, output_docx)

if __name__ == "__main__":
    success = main()
    if success:
        print("文档创建成功!")
    else:
        print("文档创建失败，请检查环境和依赖")