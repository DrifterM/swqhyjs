import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# 设置默认样式
style = doc.styles['Normal']
style.font.name = '宋体'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(12)

# 第一行：附件2 (黑体三号)
attachment = doc.add_paragraph()
run = attachment.add_run("附件2")
run.bold = True
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(16)  # 三号 ≈ 16pt
attachment.alignment = WD_ALIGN_PARAGRAPH.LEFT

# 第二行：主标题 (黑体小二)
title = doc.add_paragraph()
run = title.add_run("期货交易咨询业务档案卷内文件目录")
run.bold = True
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(18)  # 小二 ≈ 18pt
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 空一行
doc.add_paragraph()

# 创建单个统一表格 (8列)
table = doc.add_table(rows=0, cols=8)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# --- 第1行：档号 ---
row_cells = table.add_row().cells

# A1: "档号" 字段
row_cells[0].text = "档号"
row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
row_cells[0].paragraphs[0].runs[0].bold = True
row_cells[0].width = Inches(0.5)

# B1:H1: 档号内容 (合并)
merged_cell = row_cells[1].merge(row_cells[7])
merged_cell.text = "SWQH-YW·2025-D30-3-1009-0001"
merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐 ✓

# --- 第2行：案卷题名 ---
row_cells = table.add_row().cells

# A2: "案卷题名" 字段
row_cells[0].text = "案卷题名"
row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
row_cells[0].paragraphs[0].runs[0].bold = True
row_cells[0].width = Inches(0.5)

# B2:H2: 案卷题名内容 (合并)
merged_cell = row_cells[1].merge(row_cells[7])
merged_cell.text = "毛喜瑞交易咨询项目"
merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐 ✓

# --- 第3行：客户名称 + 客户代码 ---
row_cells = table.add_row().cells

# A3: "客户名称" 字段
row_cells[0].text = "客户名称"
row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
row_cells[0].paragraphs[0].runs[0].bold = True
row_cells[0].width = Inches(0.5)

# B3:D3: 客户名称值 (合并)
merged_cell = row_cells[1].merge(row_cells[3])
merged_cell.text = "毛喜瑞交易咨询项目"
merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐 ✓

# E3: "客户代码" 字段
row_cells[4].text = "客户代码"
row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
row_cells[4].paragraphs[0].runs[0].bold = True
row_cells[4].width = Inches(0.7)

# F3:H3: 客户代码值 (合并)
merged_cell = row_cells[5].merge(row_cells[7])
merged_cell.text = "8005108018"
merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐 ✓

# --- 第4行：表头 ---
header_cells = table.add_row().cells
headers = ["序号", "文件题名", "责任者", "经办人", "日期", "材料性质", "起止页号", "备注"]
for i, header in enumerate(headers):
    header_cells[i].text = header
    header_cells[i].paragraphs[0].runs[0].bold = True
    header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# 设置列宽
col_widths = [
    Inches(0.5),  # 序号
    Inches(3.0),  # 文件题名
    Inches(0.6),  # 责任者
    Inches(0.6),  # 经办人
    Inches(0.7),  # 日期
    Inches(0.7),  # 材料性质
    Inches(0.6),  # 起止页号
    Inches(0.8)   # 备注
]

for i, width in enumerate(col_widths):
    header_cells[i].width = width

# 文件题名列左对齐
header_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

# --- 文件清单 ---
files = [
    {"序号": "0001", "文件题名": "1、申报OA-毛喜瑞.pdf", "日期": "20251110"},
    {"序号": "0002", "文件题名": "2、业务（含产品、服务）洗钱风险评估表-毛喜瑞交易咨询项目.xlsx", "日期": "20251110"},
    {"序号": "0003", "文件题名": "3、合规审核意见截图.png", "日期": "20251111"},
    {"序号": "0004", "文件题名": "4、交易咨询部-魏子骥-期货交易咨询业务风险等级评估表.pdf", "日期": "20251114"},
    {"序号": "0005", "文件题名": "5、毛喜瑞伟交易咨询项目评审表决票.pdf", "日期": "20251110"},
    {"序号": "0006", "文件题名": "6、期货交易咨询业务内部评审小组立项讨论会纪要-毛喜瑞交易咨询项目.pdf", "日期": "20251114"},
    {"序号": "0007", "文件题名": "7、用印OA-毛喜瑞与申银万国期货有限公司之期货交易咨询协议.pdf", "日期": "20260127"},
    {"序号": "0008", "文件题名": "8、毛喜瑞与申银万国期货有限公司之期货交易咨询协议.pdf", "日期": "20251231"},
]

for file_item in files:
    row_cells = table.add_row().cells
    row_cells[0].text = file_item['序号']
    row_cells[1].text = file_item['文件题名']
    row_cells[4].text = file_item['日期']
    
    # 对齐
    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

# --- 新增：小计行 ---
row_cells = table.add_row().cells
row_cells[1].text = "小计"
row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
row_cells[1].paragraphs[0].runs[0].bold = True

# 保存文件
output_path = r"D:\\申银万国期货\\3.档案管理\\2、业务档案\\2. 未归档项目-2025\\毛喜瑞交易咨询项目\\卷内文件目录_居中对齐版.docx"
doc.save(output_path)
print(f"✅ 居中对齐版已成功生成: {output_path}")