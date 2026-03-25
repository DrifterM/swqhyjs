"""
真正最终完美版 - 毛喜瑞交易咨询项目
重构表格结构，确保每一列宽度精确控制
"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def create_final_catalog():
    # 项目信息
    project_name = "毛喜瑞交易咨询项目"
    client_code = "8005108018"
    product_name = "毛喜瑞交易咨询项目"
    year = "2025"
    project_num = "0001"
    archive_code = f"SWQH-YW·{year}-D30-3-1009-{project_num}"

    # 文件列表（8个以数字开头的文件）
    files = [
        {"序号": "0001", "文件题名": "1、申报OA-毛喜瑞.pdf", "日期": "20251110"},
        {"序号": "0002", "文件题名": "2、业务（含产品、服务）洗钱风险评估表-毛喜瑞交易咨询项目.xlsx", "日期": "20251110"},
        {"序号": "0003", "文件题名": "3、合规审核意见截图.png", "日期": "20251111"},
        {"序号": "0004", "文件题名": "4、交易咨询部-魏子骥-期货交易咨询业务风险等级评估表.pdf", "日期": "20251114"},
        {"序号": "0005", "文件题名": "5、毛喜瑞伟交易咨询项目评审表决票.pdf", "日期": "20251110"},
        {"序号": "0006", "文件题名": "6、期货交易咨询业务内部评审小组立项讨论会纪要-毛喜瑞交易咨询项目.pdf", "日期": "20251114"},
        {"序号": "0007", "文件题名": "7、用印OA-毛喜瑞与申银万国期货有限公司之期货交易咨询协议.pdf", "日期": "20260127"},
        {"序号": "0008", "文件题名": "8、毛喜瑞与申银万国期货有限公司之期货交易咨询协议.pdf", "日期": "20251231"},
    ]

    try:
        # 创建文档
        doc = Document()

        # 设置默认样式
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(12)  # 小四

        # 第一行：附件2 (黑体三号)
        attachment = doc.add_paragraph()
        run = attachment.add_run("附件2")
        run.bold = True
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(16)  # 三号 ≈ 16pt
        attachment.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 第二行：主标题 (黑体小二)
        title = doc.add_paragraph()
        run = title.add_run("期货交易咨询业务档案卷内文件目录")
        run.bold = True
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(18)  # 小二 ≈ 18pt
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 空一行
        doc.add_paragraph()

        # 创建单个统一表格
        table = doc.add_table(rows=0, cols=8)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # --- 表头行 ---
        header_cells = table.add_row().cells
        headers = ["档号", "", "", "", "案卷题名", "", "", ""]
        for i, header in enumerate(headers):
            if header:
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
                header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- 档号内容行 ---
        row_cells = table.add_row().cells
        row_cells[0].text = ""  # 预留位置
        row_cells[1].merge(row_cells[3])  # 合并2-4列为档号内容
        row_cells[1].text = archive_code
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        row_cells[1].paragraphs[0].paragraph_format.left_indent = Inches(0.2)

        # --- 案卷题名内容行 ---
        row_cells = table.add_row().cells
        row_cells[4].text = ""  # 预留位置
        row_cells[5].merge(row_cells[7])  # 合并6-8列为案卷题名内容
        row_cells[5].text = product_name
        row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        row_cells[5].paragraphs[0].paragraph_format.left_indent = Inches(0.2)

        # --- 客户信息行 ---
        row_cells = table.add_row().cells
        
        # 客户名称字段
        row_cells[0].text = "客户名称"
        row_cells[0].paragraphs[0].runs[0].bold = True
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 客户名称值
        row_cells[1].merge(row_cells[3])
        row_cells[1].text = project_name
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        row_cells[1].paragraphs[0].paragraph_format.left_indent = Inches(0.2)
        
        # 客户代码字段
        row_cells[4].text = "客户代码"
        row_cells[4].paragraphs[0].runs[0].bold = True
        row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 客户代码值
        row_cells[5].merge(row_cells[7])
        row_cells[5].text = client_code
        row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        row_cells[5].paragraphs[0].paragraph_format.left_indent = Inches(0.2)

        # --- 实际表头行 ---
        header_cells = table.add_row().cells
        headers = ["序号", "文件题名", "责任者", "经办人", "日期", "材料性质", "起止页号", "备注"]
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 设置列宽
        col_widths = [
            Inches(0.5),  # 序号
            Inches(3.0),  # 文件题名
            Inches(0.6),  # 责任者
            Inches(0.6),  # 经办人
            Inches(0.7),  # 日期
            Inches(0.7),  # 材料性质
            Inches(0.6),  # 起止页号
            Inches(0.8)   # 备注
        ]
        
        for i, width in enumerate(col_widths):
            header_cells[i].width = width

        # 文件题名列左对齐
        header_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 添加文件行
        for file_item in files:
            row_cells = table.add_row().cells
            row_cells[0].text = file_item['序号']
            row_cells[1].text = file_item['文件题名']
            row_cells[4].text = file_item['日期']

            # 对齐
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 不添加页脚（移除生成时间）

        # 保存文件
        output_path = r"D:\\申银万国期货\\3.档案管理\\2、业务档案\\2. 未归档项目-2025\\毛喜瑞交易咨询项目\\卷内文件目录_真正最终完美版.docx"
        doc.save(output_path)

        print(f"✅ 真正最终完美版已成功生成: {output_path}")
        print(f"文件大小: {os.path.getsize(output_path)} 字节")

    except Exception as e:
        print(f"❌ 创建Word文档时出错: {type(e).__name__}")
        print(f"错误详情: {e}")

if __name__ == "__main__":
    create_final_catalog()