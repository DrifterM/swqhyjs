"""
创建最终修复版卷内文件目录
- 档号、案卷题名、客户信息均分字段显示
- 文件序号为四位数 (0001)
- 优化表格列宽
- 移除生成时间
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

# 项目信息
customer_name = "毛喜瑞交易咨询项目"
code = "8005108018"
product_name = "毛喜瑞交易咨询项目"
year = "2025"
project_num = "0001"  # 四位编号
archive_code = f"SWQH-YW·{year}-D30-3-1009-{project_num}"

# 文件列表
file_list = [
    {"序号": "0001", "文件题名": "1、申报OA-毛喜瑞.pdf", "日期": "20251110"},
    {"序号": "0002", "文件题名": "2、业务（含产品、服务）洗钱风险评估表-毛喜瑞交易咨询项目.xlsx", "日期": "20251110"},
    {"序号": "0003", "文件题名": "3、合规审核意见截图.png", "日期": "20251111"},
    {"序号": "0004", "文件题名": "4、交易咨询部-魏子骥-期货交易咨询业务风险等级评估表.pdf", "日期": "20251114"},
    {"序号": "0005", "文件题名": "5、毛喜瑞伟交易咨询项目评审表决票.pdf", "日期": "20251110"},
    {"序号": "0006", "文件题名": "6、期货交易咨询业务内部评审小组立项讨论会纪要-毛喜瑞交易咨询项目.pdf", "日期": "20251114"},
    {"序号": "0007", "文件题名": "7、用印OA-毛喜瑞与申银万国期货有限公司之期货交易咨询协议.pdf", "日期": "20260127"},
    {"序号": "0008", "文件题名": "8、毛喜瑞与申银万国期货有限公司之期货交易咨询协议.pdf", "日期": "20251231"},
]

try:
    # 创建文档
doc = Document()
    
    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)  # 小四
    
    # 第一行：附件2 (黑体三号)
    attachment = doc.add_paragraph()
    run = attachment.add_run("附件2")
    run.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(16)  # 三号 ≈ 16pt
    attachment.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 第二行：主标题 (黑体小二)
    title = doc.add_paragraph()
    run = title.add_run("期货交易咨询业务档案卷内文件目录")
    run.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(18)  # 小二 ≈ 18pt
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 空一行
    doc.add_paragraph()
    
    # 创建单个统一表格
    table = doc.add_table(rows=0, cols=8)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # --- 第一行：档号 (两列) ---
    row_cells = table.add_row().cells
    row_cells[0].width = Inches(1.0)     # 左侧字段列
    row_cells[1].width = Inches(6.0)    # 右侧内容列
    row_cells[0].merge(row_cells[3])   # 合并前4列作为"档号"字段
    row_cells[4].merge(row_cells[7])   # 合并后4列作为内容
    
    row_cells[0].text = "档号"
    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row_cells[0].paragraphs[0].runs[0].bold = True
    
    row_cells[4].text = archive_code
    row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    row_cells[4].paragraphs[0].paragraph_format.left_indent = Inches(0.2)
    
    # --- 第二行：案卷题名 (两列) ---
    row_cells = table.add_row().cells
    row_cells[0].merge(row_cells[3])
    row_cells[4].merge(row_cells[7])
    
    row_cells[0].text = "案卷题名"
    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row_cells[0].paragraphs[0].runs[0].bold = True
    
    row_cells[4].text = product_name
    row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    row_cells[4].paragraphs[0].paragraph_format.left_indent = Inches(0.2)
    
    # --- 第三行：客户名称 + 客户代码 (四列) ---
    row_cells = table.add_row().cells
    
    # 客户名称 (两列)
    row_cells[0].merge(row_cells[1])
    row_cells[0].text = "客户名称"
    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row_cells[0].paragraphs[0].runs[0].bold = True
    
    row_cells[2].merge(row_cells[3])
    row_cells[2].text = customer_name
    row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    row_cells[2].paragraphs[0].paragraph_format.left_indent = Inches(0.2)
    
    # 客户代码 (两列)
    row_cells[4].merge(row_cells[5])
    row_cells[4].text = "客户代码"
    row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row_cells[4].paragraphs[0].runs[0].bold = True
    
    row_cells[6].merge(row_cells[7])
    row_cells[6].text = code
    row_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    row_cells[6].paragraphs[0].paragraph_format.left_indent = Inches(0.2)
    
    # --- 第四行：表头 ---
    header_cells = table.add_row().cells
    headers = ["序号", "文件题名", "责任者", "经办人", "日期", "材料性质", "起止页号", "备注"]
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 设置列宽
    # 总宽度约 7 英寸
    header_cells[0].width = Inches(0.5)   # 序号
    header_cells[1].width = Inches(3.0)   # 文件题名 (加宽)
    header_cells[2].width = Inches(0.6)   # 责任者
    header_cells[3].width = Inches(0.6)   # 经办人
    header_cells[4].width = Inches(0.7)   # 日期
    header_cells[5].width = Inches(0.7)   # 材料性质
    header_cells[6].width = Inches(0.6)   # 起止页号
    header_cells[7].width = Inches(0.8)   # 备注
    
    # 文件题名列左对齐
    header_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 添加文件行
    for file_item in file_list:
        row_cells = table.add_row().cells
        row_cells[0].text = file_item['序号']
        row_cells[1].text = file_item['文件题名']
        row_cells[4].text = file_item['日期']
        
        # 对齐
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 保存文件
    output_path = r"D:\\申银万国期货\\3.档案管理\\2、业务档案\\2. 未归档项目-2025\\毛喜瑞交易咨询项目\\卷内文件目录.docx"
    doc.save(output_path)
    
    print(f"✅ 最终修复版已成功生成: {output_path}")
    print(f"文件大小: {os.path.getsize(output_path)} 字节")
    
except Exception as e:
    print(f"❌ 创建Word文档时出错: {type(e).__name__}")
    print(f"错误详情: {e}")