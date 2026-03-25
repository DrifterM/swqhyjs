from docx import Document

doc = Document('D:\\申银万国期货\\3.档案管理\\2、业务档案\\2. 未归档项目-2025\\毛喜瑞交易咨询项目\\卷内文件目录.docx')

# Read all paragraphs
texts = []
for para in doc.paragraphs:
    texts.append(para.text)

# Read all tables
tables_data = []
for table in doc.tables:
    table_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text.strip())
        table_data.append(row_data)
    tables_data.append(table_data)

# Output
{'paragraphs': texts, 'tables': tables_data}