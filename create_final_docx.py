"""
创建符合正式规范的卷内文件目录Word文档
使用正确的Python环境：D:\\CoPaw\\venv\\Scripts\\python.exe
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

# 项目信息
PROJECT_NAME = "毛喜瑞交易咨询项目"
CLIENT_CODE = "8005108018"
PRODUCT_NAME = "毛喜瑞交易咨询项目"
YEAR = "2025"
PROJECT_NUM = "0001"
ARCHIVE_CODE = f"SWQH-YW·{YEAR}-D30-3-1009-{PROJECT_NUM}"

# 文件列表
file_list = [
    {"序号": "1", "文件题名": "1、申报OA-毛喜瑞.pdf", "日期": "20251110"},
    {"序号": "2", "文件题名": "2、业务（含产品、服务）洗钱风险评估表-毛喜瑞交易咨询项目.xlsx", "日期": "20251110"},
    {"序号": "3", "文件题名": "3、合规审核意见截图.png", "日期": "20251111"},
    {"序号": "4", "文件题名": "4、交易咨询部-魏子骥-期货交易咨询业务风险等级评估表.pdf", "日期": "20251114"},
    {"序号": "5", "文件题名": "5、毛喜瑞伟交易咨询项目评审表决票.pdf", "日期": "20251110"},
    {"序号": "6", "文件题名": "6、期货交易咨询业务内部评审小组立项讨论会纪要-毛喜瑞交易咨询项目.pdf", "日期": "20251114"},
    {"序号": "7", "文件题名": "7、用印OA-毛喜瑞与申银万国期货有限公司之期货交易咨询协议.pdf", "日期": "20260127"},
    {"序号": "8", "文件题名": "8、毛喜瑞与申银万国期货有限公司之期货交易咨询协议.pdf", "日期": "20251231"},
]

try:
    # 创建新文档
    doc = Document()
    
    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)  # 小四
    
    # 第一行：附件2 (黑体三号)
    attachment = doc.add_paragraph()
    run = attachment.add_run("附件2")
    run.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(16)  # 三号 ≈ 16pt
    attachment.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 第二行：主标题 (黑体小二)
    title = doc.add_paragraph()
    run = title.add_run("期货交易咨询业务档案卷内文件目录")
    run.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(18)  # 小二 ≈ 18pt
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 空一行
    doc.add_paragraph()
    
    # 创建单个统一表格
    table = doc.add_table(rows=0, cols=8)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 第一行：档号
    row_cells = table.add_row().cells
    row_cells[0].merge(row_cells[-1])  # 合并所有列
    row_cells[0].text = ARCHIVE_CODE
    row_cells[0].paragraphs[0].runs[0].bold = True
    row_cells[0].paragraphs[0].runs[0].font.size = Pt(13)  # 稍大字号
    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 第二行：案卷题名
    row_cells = table.add_row().cells
    row_cells[0].merge(row_cells[-1])
    row_cells[0].text = PRODUCT_NAME
    row_cells[0].paragraphs[0].runs[0].bold = True
    row_cells[0].paragraphs[0].runs[0].font.size = Pt(13)
    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 第三行：客户名称 + 客户代码
    row_cells = table.add_row().cells
    row_cells[0].merge(row_cells[3])  # 合并前4列
    row_cells[0].text = f"客户名称：{PROJECT_NAME}"
    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    row_cells[0].paragraphs[0].runs[0].font.size = Pt(12)
    
    row_cells[4].merge(row_cells[7])  # 合并后4列
    row_cells[4].text = f"客户代码：{CLIENT_CODE}"
    row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    row_cells[4].paragraphs[0].runs[0].font.size = Pt(12)
    
    # 第四行：表头
    header_cells = table.add_row().cells
    headers = ["序号", "文件题名", "责任者", "经办人", "日期", "材料性质", "起止页号", "备注"]
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 设置表头宽度（通过段落缩进模拟）
    # 文件题名列加宽
    header_cells[1].paragraphs[0].paragraph_format.left_indent = Inches(0.5)
    
    # 添加文件行
    for file_item in file_list:
        row_cells = table.add_row().cells
        row_cells[0].text = file_item['序号']
        row_cells[1].text = file_item['文件题名']
        row_cells[4].text = file_item['日期']
        
        # 对齐
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 文件题名列左对齐
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 添加页脚信息
    footer = doc.add_paragraph()
    footer.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 保存文件
    output_path = r"D:\\申银万国期货\\3.档案管理\\2、业务档案\\2. 未归档项目-2025\\毛喜瑞交易咨询项目\\卷内文件目录.docx"
    doc.save(output_path)
    
    print(f"✅ 最终版Word文档已成功生成: {output_path}")
    print(f"文件大小: {os.path.getsize(output_path)} 字节")
    
except Exception as e:
    print(f"❌ 创建Word文档时出错: {type(e).__name__}")
    print(f"错误详情: {e}")