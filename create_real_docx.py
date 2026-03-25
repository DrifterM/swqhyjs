from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import datetime

# 常量
YEAR = "2025"
PROJECT_NUM = "0001"
PROJECT_NAME = "毛喜瑞交易咨询项目"
CLIENT_CODE = "8005108018"
PRODUCT_NAME = "毛喜瑞交易咨询项目"
ARCHIVE_CODE = f"SWQH-YW·{YEAR}-D30-3-1009-{PROJECT_NUM}"

# 文件列表
file_list = [
    {"序号": "0001", "文件题名": "1、申报OA-毛喜瑞.pdf", "日期": "20251110"},
    {"序号": "0002", "文件题名": "2、业务（含产品、服务）洗钱风险评估表-毛喜瑞交易咨询项目.xlsx", "日期": "20251110"},
    {"序号": "0003", "文件题名": "3、合规审核意见截图.png", "日期": "20251111"},
    {"序号": "0004", "文件题名": "4、交易咨询部-魏子骥-期货交易咨询业务风险等级评估表.pdf", "日期": "20251114"},
    {"序号": "0005", "文件题名": "5、毛喜瑞伟交易咨询项目评审表决票.pdf", "日期": "20251114"},
    {"序号": "0006", "文件题名": "6、期货交易咨询业务内部评审小组立项讨论会纪要-毛喜瑞交易咨询项目.pdf", "日期": "20251114"},
    {"序号": "0007", "文件题名": "7、用印OA-毛喜瑞与申银万国期货有限公司之期货交易咨询协议.pdf", "日期": "20260127"},
    {"序号": "0008", "文件题名": "8、毛喜瑞与申银万国期货有限公司之期货交易咨询协议.pdf", "日期": "20251231"},
]

try:
    # 创建新文档
    doc = Document()
    
    # 设置文档样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    
    # 标题
    title = doc.add_heading('卷内文件目录', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息表格 (2行4列)
    info_table = doc.add_table(rows=2, cols=4)
    info_table.style = 'Table Grid'
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 第一行
    info_table.cell(0, 0).text = "客户名称"
    info_table.cell(0, 1).text = PROJECT_NAME
    info_table.cell(0, 2).text = "档号"
    info_table.cell(0, 3).text = ARCHIVE_CODE
    
    # 第二行
    info_table.cell(1, 0).text = "案卷题名"
    info_table.cell(1, 1).text = PRODUCT_NAME
    info_table.cell(1, 2).text = "客户代码"
    info_table.cell(1, 3).text = CLIENT_CODE
    
    # 设置表格单元格样式
    for row in info_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 空行
    doc.add_paragraph()
    
    # 文件列表表格
    file_table = doc.add_table(rows=1, cols=7)
    file_table.style = 'Table Grid'
    file_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    headers = ["序号", "文件题名", "日期", "责任者", "经办人", "起止页号", "备注"]
    header_cells = file_table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加文件行
    for file_item in file_list:
        row_cells = file_table.add_row().cells
        row_cells[0].text = file_item['序号']
        row_cells[1].text = file_item['文件题名']
        row_cells[2].text = file_item['日期']
        row_cells[3].text = ""
        row_cells[4].text = ""
        row_cells[5].text = ""
        row_cells[6].text = ""
        
        # 对齐
        for i in [0, 2, 3, 4, 5, 6]:
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 页脚信息
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 保存文件
    output_path = r"D:\申银万国期货\3.档案管理\2、业务档案\2. 未归档项目-2025\毛喜瑞交易咨询项目\卷内文件目录_标准版.docx"
    doc.save(output_path)
    
    print(f"✅ Word文档已成功生成: {output_path}")
    print(f"文件大小: {os.path.getsize(output_path)} 字节")
    
except Exception as e:
    print(f"❌ 创建Word文档时出错: {type(e).__name__}")
    print(f"错误详情: {e}")
    print("\n尝试生成简化版本...")
    
    # 生成一个基本的文本文件作为备份
    backup_content = f"""卷内文件目录

客户名称：{PROJECT_NAME}
档号：{ARCHIVE_CODE}
案卷题名：{PRODUCT_NAME}
客户代码：{CLIENT_CODE}

序号  文件题名                             日期      责任者 经办人 起止页号 备注
{"-"*100}"""
    
    for file_item in file_list:
        backup_content += f"\n{file_item['序号']:4}  {file_item['文件题名']:50}  {file_item['日期']}        -       -       -       -"
    
    backup_path = r"D:\申银万国期货\3.档案管理\2、业务档案\2. 未归档项目-2025\毛喜瑞交易咨询项目\卷内文件目录_备份.txt"
    with open(backup_path, 'w', encoding='utf-8') as f:
        f.write(backup_content)
    
    print(f"✅ 已生成备份文本文件: {backup_path}")