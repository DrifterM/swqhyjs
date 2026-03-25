"""
技能脚本：生成标准化卷内文件目录
"""

import os
import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import json

# 导入工具函数
try:
    from utils import find_latest_ledger, get_client_code_from_ledger
except ImportError:
    # 当前目录加入路径
    import sys
    from pathlib import Path
    sys.path.append(str(Path(__file__).parent))
    from utils import find_latest_ledger, get_client_code_from_ledger

def load_config():
    """加载配置"""
    config_path = os.path.join(os.path.dirname(__file__), "config.json")
    with open(config_path, 'r', encoding='utf-8') as f:
        return json.load(f)

def scan_files(project_path):
    """扫描项目文件夹，获取以数字开头的文件"""
    files = []
    for file in os.listdir(project_path):
        if file.startswith(("0", "1", "2", "3", "4", "5", "6", "7", "8", "9")) and not file.endswith((".docx", ".html")):
            files.append(file)
    # 按名称排序
    files.sort(key=lambda x: int(x.split('、')[0]) if x.split('、')[0].isdigit() else 0)
    return files

def generate_catalog(project_path, client_name=None, client_code=None, year="2025", seq_num="0001", auto_match_code=False, ledger_path_pattern=None):
    """生成卷内文件目录"""
    config = load_config()

    # 推断客户名称
    if not client_name:
        client_name = os.path.basename(project_path)
    
    # 如果启用自动匹配且未提供客户代码
    if auto_match_code and not client_code and ledger_path_pattern:
        print(f"🔍 正在从台账中查找 '{client_name}' 的客户代码...")
        ledger_path = find_latest_ledger(ledger_path_pattern)
        if not ledger_path:
            print(f"❌ 未找到匹配的台账文件: {ledger_path_pattern}")
        else:
            print(f"✅ 找到台账: {ledger_path}")
            code = get_client_code_from_ledger(ledger_path, client_name)
            if code:
                client_code = code
                print(f"✅ 匹配到客户代码: {client_code}")
            else:
                print(f"⚠️  在台账中未找到客户 '{client_name}' 的代码")
    
    # 创建文档
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = config['default_font']
    style._element.rPr.rFonts.set(qn('w:eastAsia'), config['default_font'])
    style.font.size = Pt(12)

    # 第一行：附件2
    attachment = doc.add_paragraph()
    run = attachment.add_run(config['template']['title1'])
    run.bold = True
    run.font.name = config['bold_font']
    run._element.rPr.rFonts.set(qn('w:eastAsia'), config['bold_font'])
    run.font.size = Pt(16)
    attachment.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 第二行：主标题
    title = doc.add_paragraph()
    run = title.add_run(config['template']['title2'])
    run.bold = True
    run.font.name = config['bold_font']
    run._element.rPr.rFonts.set(qn('w:eastAsia'), config['bold_font'])
    run.font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 空一行
doc.add_paragraph()

# 创建表格
table = doc.add_table(rows=0, cols=8)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# --- 第1行：档号 ---
row_cells = table.add_row().cells
row_cells[0].text = "档号"
row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
row_cells[0].paragraphs[0].runs[0].bold = True
row_cells[0].width = Inches(0.5)

merged_cell = row_cells[1].merge(row_cells[7])
archive_code = config['archive_code_pattern'].replace("YYYY", year).replace("XXXX", seq_num)
merged_cell.text = archive_code
merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# --- 第2行：案卷题名 ---
row_cells = table.add_row().cells
row_cells[0].text = "案卷题名"
row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
row_cells[0].paragraphs[0].runs[0].bold = True
row_cells[0].width = Inches(0.5)

merged_cell = row_cells[1].merge(row_cells[7])
merged_cell.text = client_name
merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# --- 第3行：客户信息 ---
row_cells = table.add_row().cells

row_cells[0].text = "客户名称"
row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
row_cells[0].paragraphs[0].runs[0].bold = True
row_cells[0].width = Inches(0.5)

merged_cell = row_cells[1].merge(row_cells[3])
merged_cell.text = client_name
merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

row_cells[4].text = "客户代码"
row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
row_cells[4].paragraphs[0].runs[0].bold = True
row_cells[4].width = Inches(0.7)

merged_cell = row_cells[5].merge(row_cells[7])
merged_cell.text = client_code or ""
merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# --- 第4行：表头 ---
header_cells = table.add_row().cells
for i, header in enumerate(config['template']['header_row']):
    header_cells[i].text = header
    header_cells[i].paragraphs[0].runs[0].bold = True
    header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# 设置列宽
for i, width in enumerate(config['template']['column_widths']):
    header_cells[i].width = Inches(width)

# 文件题名列左对齐
header_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

# --- 扫描并添加文件 ---
files = scan_files(project_path)
for i, filename in enumerate(files):
    row_cells = table.add_row().cells
    row_cells[0].text = f"{i+1:04d}"  # 四位序号
    row_cells[1].text = filename
    # 尝试提取修改日期（可选）
    try:
        mtime = os.path.getmtime(os.path.join(project_path, filename))
        date_str = str(int(mtime // 1e9) % 100000000)[-8:]  # 简化处理
        row_cells[4].text = date_str
    except:
        pass
    # 对齐
    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

# --- 小计行 ---
row_cells = table.add_row().cells
row_cells[1].text = "小计"
row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
row_cells[1].paragraphs[0].runs[0].bold = True

# 保存
doc.save(os.path.join(project_path, "卷内文件目录.docx"))
print(f"✅ 卷内文件目录已生成于: {project_path}")

if __name__ == "__main__":
    # 命令行调用支持
    if len(sys.argv) < 2:
        print("Usage: python generate_catalog.py <project_path> [client_name] [client_code] [year] [seq_num]")
        sys.exit(1)
    project_path = sys.argv[1]
    client_name = sys.argv[2] if len(sys.argv) > 2 else None
    client_code = sys.argv[3] if len(sys.argv) > 3 else None
    year = sys.argv[4] if len(sys.argv) > 4 else "2025"
    seq_num = sys.argv[5] if len(sys.argv) > 5 else "0001"
    
    # 确保路径为绝对路径且存在
    project_path = os.path.abspath(project_path)
    if not os.path.exists(project_path):
        print(f"Error: Project path does not exist: {project_path}")
        sys.exit(1)
    
    generate_catalog(project_path, client_name, client_code, year, seq_num)